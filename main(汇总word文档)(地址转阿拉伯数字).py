import os
import re
import fnmatch
from tqdm import tqdm
from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Font


def remove_punctuation(text):
    return re.sub(r'[^\w\s]', '', text)


def get_hot_number(cell_text):
    # 支持多种热号格式
    match = re.search(r'[(（]?[热暖]号[:：\s]*([\d]+)[)）]?', cell_text)
    if match:
        return match.group(1)
    return ''


def get_address_from_hot_number(cell_text):
    match = re.search(r'(.*?)([(（]?[热暖]号)', cell_text)
    if match:
        return match.group(1)
    return ''


def get_hot_number_from_table(table):
    for row in table.rows:
        cells = row.cells
        for i, cell in enumerate(cells):
            if "申请单位" in cell.text and i + 1 < len(cells):
                return get_hot_number(cells[i + 1].text)
    return ''


def chinese_to_arabic(cn_num):
    cn_units = {'十': 10, '百': 100, '千': 1000, '万': 10000}
    cn_nums = {'零': 0, '一': 1, '二': 2, '三': 3, '四': 4, '五': 5, '六': 6, '七': 7, '八': 8, '九': 9}
    arabic_num = 0
    unit = 1

    for char in reversed(cn_num):
        if char in cn_units:
            unit = cn_units[char]
        else:
            num = cn_nums[char]
            arabic_num += num * unit
            unit = 1

    return arabic_num


def convert_chinese_to_arabic_numbers(text):
    match = re.findall(r'[零一二三四五六七八九十百千万]+', text)
    if match:
        for cn_num in match:
            arabic_num = chinese_to_arabic(cn_num)
            text = text.replace(cn_num, str(arabic_num))
    return text


def extract_info_from_cell(cell_text):
    handler_match = re.search(r'受理人[:：\s]*(.*)', cell_text)
    handler = handler_match.group(1) if handler_match else ""

    date_match = re.search(r'(\d{4}\s?年\s?\d{1,2}\s?月\s?\d{1,2}\s?日)', cell_text)
    date = date_match.group(1).replace("年", ".").replace("月", ".").replace("日", "").replace(" ",
                                                                                               "") if date_match else ""

    address = get_address_from_hot_number(cell_text)
    address = remove_punctuation(address.strip())
    address = convert_chinese_to_arabic_numbers(address)

    return address, handler, date


def main():
    wb = Workbook()
    ws = wb.active

    titles = ['热号', '地址', '联系人', '联系电话', '编号', '文件名', '文件路径', '备注', '受理人',
              '受理时间']  # 添加了新的列标题
    for col, title in enumerate(titles, start=1):
        cell = ws.cell(row=1, column=col)
        cell.font = Font(bold=True)
        cell.value = title

    folder_path = "C:/Users/Administrator/Desktop/农村非低保残疾人/自己汇总"
    files = []
    for root, dirs, filenames in os.walk(folder_path):
        for filename in fnmatch.filter(filenames, "*.doc*"):
            files.append(os.path.join(root, filename))

    current_row = 2
    total_files = len(files)
    for idx, file in enumerate(tqdm(files), start=1):
        ws.cell(row=current_row, column=6).value = os.path.basename(file)  # 先记录文件名

        try:
            doc = Document(file)
        except Exception as e:  # 捕获错误，跳过无法读取的文档
            ws.cell(row=current_row, column=8).value = f"错误: {str(e)}"  # 将错误信息记录在备注列
            current_row += 1
            continue

        file_name = os.path.basename(file)  ### 获取文件名
        relative_path = os.path.relpath(file, folder_path)  ### 获取相对路径
        path_parts = relative_path.split(os.path.sep)[:-1]  ### 分割路径并移除文件名部分
        formatted_path = '-'.join(path_parts)  ### 格式化路径为所需格式

        for table in doc.tables:
            contact_name = ''
            contact_phone = ''
            hot_number = get_hot_number_from_table(table)
            ws.cell(row=current_row, column=1).value = hot_number

            for row in table.rows:
                cells = row.cells
                for i, cell in enumerate(cells):
                    if "申请单位" in cell.text and i + 1 < len(cells):
                        address, handler, date = extract_info_from_cell(cells[i + 1].text)
                        ws.cell(row=current_row, column=2).value = address
                        ws.cell(row=current_row, column=9).value = handler
                        ws.cell(row=current_row, column=10).value = date



                    elif "联系人" in cell.text and i + 1 < len(cells):
                        contact_name = cells[i + 1].text.replace(" ", "")


                    elif "联系电话" in cell.text and i + 1 < len(cells):
                        contact_phone = cells[i + 1].text.replace(" ", "")

            ws.cell(row=current_row, column=3).value = contact_name
            ws.cell(row=current_row, column=4).value = contact_phone
            ws.cell(row=current_row, column=6).value = file_name  ### 记录文件名
            ws.cell(row=current_row, column=7).value = formatted_path  ### 记录文件路径
            current_row += 1

        for para in doc.paragraphs:
            match = re.search(r'编号[：:\s]*([\d-]+)', para.text)
            if match:
                ws.cell(row=(current_row - 1), column=5).value = match.group(1)

    output_filename = os.path.basename(folder_path) + ".xlsx"
    output_file = os.path.join(folder_path, output_filename)
    wb.save(output_file)
    print("处理完成")

if __name__ == "__main__":
    main()

