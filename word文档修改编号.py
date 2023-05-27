import os
import re
import fnmatch
import docx
from tqdm import tqdm
from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Font


def remove_punctuation(text):
    return re.sub(r'[^\w\s]', '', text)


def get_hot_number(cell_text):
    match = re.search(r'[(（]?[热暖]号[:：\s]*([\d]+)[)）]?', cell_text)
    if match:
        return match.group(1)
    return ''


def get_hot_number_from_table(table):
    for row in table.rows:
        cells = row.cells
        for i, cell in enumerate(cells):
            if "申请单位" in cell.text and i + 1 < len(cells):
                return get_hot_number(cells[i + 1].text)
    return ''


def process_text(text):
    match = re.search(r'具体事项及内容：(.*?)(?=户)', text)
    if match:
        result = match.group(1).replace('用', '')
        result = remove_punctuation(result.strip())
        return result
    return ''


def extract_info_from_cell(cell_text):
    handler_match = re.search(r'受理人[:：\s]*(.*)', cell_text)
    handler = handler_match.group(1) if handler_match else ""

    date_match = re.search(r'(\d{4}\s?年\s?\d{1,2}\s?月\s?\d{1,2}\s?日)', cell_text)
    if date_match:
        date = date_match.group(1).replace("年", "").replace("月", "").replace("日", "").replace(" ", "")
    else:
        date = ""

    return process_text(cell_text), handler, date


def get_folder_display_name(folder_path):
    main_folder_name = os.path.basename(os.path.dirname(folder_path))
    sub_folder_name = os.path.basename(os.path.dirname(os.path.dirname(folder_path)))
    return f"{sub_folder_name}-{main_folder_name}"


def set_word_format(text_run):
    text_run.font.name = "黑体"
    text_run._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), "黑体")
    text_run.font.size = docx.shared.Pt(11)
    text_run.font.bold = True


def main():
    wb = Workbook()
    ws = wb.active

    titles = ['文件名', '原编号', '新编号', '表格编号', '文件路径', '备注', '备注2']
    for col, title in enumerate(titles, start=1):
        cell = ws.cell(row=1, column=col)
        cell.font = Font(bold=True)
        cell.value = title

    folder_path = "C:/Users/Administrator/Desktop/农村非低保残疾人/自己汇总"
    files = []
    for root, dirs, filenames in os.walk(folder_path):
        for filename in fnmatch.filter(filenames, "*.doc*"):
            files.append(os.path.join(root, filename))

    current_row = 2
    total_files = len(files)
    for idx, file in enumerate(tqdm(files), start=1):
        ws.cell(row=current_row, column=1).value = os.path.basename(file)  # 记录文件名
        ws.cell(row=current_row, column=5).value = get_folder_display_name(file)  # 记录文件路径

        try:
            doc = Document(file)
        except Exception as e:  # 捕获错误，跳过无法读取的文档
            ws.cell(row=current_row, column=7).value = f"错误: {str(e)}"  # 将错误信息记录在备注2列
            current_row += 1
            continue

        for table in doc.tables:
            hot_number = get_hot_number_from_table(table)

            for row in table.rows:
                cells = row.cells
                for i, cell in enumerate(cells):
                    if "具体事项及内容：" in cell.text:
                        _, handler, date = extract_info_from_cell(cell.text)
                        new_number = f"{date}-{hot_number}"

        for para_idx, para in enumerate(doc.paragraphs):
            match = re.search(r'编号[：:\s]*([\d-]+)', para.text)
            if match:
                old_number = match.group(1)
                ws.cell(row=current_row, column=2).value = old_number
                if old_number == new_number:
                    ws.cell(row=current_row, column=4).value = old_number
                    ws.cell(row=current_row, column=6).value = "无变化"
                else:
                    doc.paragraphs[para_idx].text = para.text.replace(old_number, new_number)
                    ws.cell(row=current_row, column=3).value = new_number
                    ws.cell(row=current_row, column=6).value = "修改"

                    set_word_format(para.runs[0])

        doc.save(file)  # 保存更改后的Word文件
        current_row += 1

    output_file = os.path.join(folder_path, "output.xlsx")
    wb.save(output_file)
    print("处理完成")


if __name__ == "__main__":
    main()
